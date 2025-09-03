import re
import json
import time
import pdfplumber
from pathlib import Path
from docx import Document
from docx2pdf import convert

OUTPUT_DIR = Path("extracted")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

NORM_PATTERN = re.compile(r'^\d+(\.\d+)*')

def process_norm_document(file_path: Path):
    source = file_path.stem
    text_norms = []
    class_norms = []

    # === Загрузка уже существующих
    text_file = OUTPUT_DIR / "temp_norms.json"
    class_file = OUTPUT_DIR / "temp_class_norms.json"
    existing_text = json.loads(text_file.read_text(encoding="utf-8")) if text_file.exists() else []
    existing_class = json.loads(class_file.read_text(encoding="utf-8")) if class_file.exists() else []

    existing_text_ids = {n["full_id"] for n in existing_text}
    existing_class_ids = {n["full_id"] for n in existing_class}
    processed_sources = {n["source"] for n in existing_text + existing_class if "source" in n}

    if source in processed_sources:
        print(f"⏭️ Документ уже обработан: {source}")
        return [], []  # <- Возвращаем пустые списки

    # === Извлечение текстовых норм
    doc = Document(file_path)
    current = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = NORM_PATTERN.match(text)
        if match:
            if current:
                text_norms.append(current)
            current = {"id": match.group(), "text": text}
        elif current:
            current["text"] += " " + text
    if current:
        text_norms.append(current)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines = cell.text.strip().split("\n")
                for line in lines:
                    text = line.strip()
                    if not text:
                        continue
                    match = NORM_PATTERN.match(text)
                    if match:
                        if current:
                            text_norms.append(current)
                        current = {"id": match.group(), "text": text}
                    elif current:
                        current["text"] += " " + text
                if current:
                    text_norms.append(current)
                    current = None

    for norm in text_norms:
        norm["source"] = source
        norm["full_id"] = f"{source}:{norm['id']}"
        if norm["full_id"] not in existing_text_ids:
            existing_text.append(norm)
            existing_text_ids.add(norm["full_id"])

    # === Конвертация DOCX → PDF и парсинг таблиц
    pdf_path = file_path.with_suffix(".pdf")
    try:
        convert(str(file_path), str(pdf_path))
    except Exception as e:
        print(f"❌ Ошибка конвертации в PDF: {e}")
        return text_norms, []  # <- возвращаем что уже смогли получить

    previous_indicator = None
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue
                header_row_idx = 1 if any(c is None for c in table[0]) else 0
                raw_headers = [str(h).strip().replace("\n", " ") for h in table[header_row_idx]]
                if header_row_idx == 1:
                    first_row = table[0]
                    headers = [f"{first_row[i]} {raw_headers[i]}".strip() if first_row[i] else raw_headers[i]
                               for i in range(len(raw_headers))]
                else:
                    headers = raw_headers
                for row in table[header_row_idx + 1:]:
                    row = [str(c).strip().replace("\n", " ") if c else "" for c in row]
                    indicator = row[0] or previous_indicator
                    previous_indicator = indicator
                    values = {headers[i]: row[i] for i in range(1, min(len(headers), len(row))) if headers[i] and row[i]}
                    if values:
                        full_id = f"{source}:{indicator}"
                        if full_id not in existing_class_ids:
                            entry = {
                                "indicator": indicator,
                                "values": values,
                                "source": source,
                                "full_id": full_id
                            }
                            existing_class.append(entry)
                            class_norms.append(entry)  # сохраняем для возврата
                            existing_class_ids.add(full_id)

    # === Сохранение
    with open(text_file, "w", encoding="utf-8") as f:
        json.dump(existing_text, f, ensure_ascii=False, indent=2)
    with open(class_file, "w", encoding="utf-8") as f:
        json.dump(existing_class, f, ensure_ascii=False, indent=2)

    print(f"✅ Завершено: {file_path.name}")
    print(f"📝 Всего текстовых норм: {len(existing_text)}")
    print(f"📊 Всего табличных норм: {len(existing_class)}")

    return text_norms, class_norms  # ✅ Возвращаем списки
